import os
import pickle
import glob
import re
import shutil
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument  # python-docx for Word files
from langchain.schema import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

# Function: Clean text
def clean_text(text):
    text = re.sub(r"[^a-zA-Z0-9.,_!?'\s]", "", text)  # Remove unwanted symbols
    text = re.sub(r"\s+", " ", text)  # Normalize spaces
    return text.strip()

# Function: Extract text from PDFs
def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join([page.get_text("text") for page in doc])
        return clean_text(text)
    except Exception as e:
        print(f"❌ Error extracting text from PDF: {pdf_path} - {e}")
        return ""

# Function: Extract text from Word documents
def extract_text_from_docx(docx_path):
    try:
        doc = DocxDocument(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return clean_text(text)
    except Exception as e:
        print(f"❌ Error extracting text from DOCX: {docx_path} - {e}")
        return ""

# Function: Preprocess Documents (PDF & DOCX only)
def preprocess_documents(pdf_files, docx_files):
    """
    Preprocesses documents for vectorization.
    - Supports only PDFs and Word Documents.
    """
    documents = []

    # Process PDFs
    for pdf_file in pdf_files:
        extracted_text = extract_text_from_pdf(pdf_file)
        if extracted_text:
            metadata = {"source": pdf_file}
            documents.append(Document(page_content=extracted_text, metadata=metadata))

    # Process Word Documents
    for docx_file in docx_files:
        extracted_text = extract_text_from_docx(docx_file)
        if extracted_text:
            metadata = {"source": docx_file}
            documents.append(Document(page_content=extracted_text, metadata=metadata))

    return documents

# Function: Reset ChromaDB only if needed
def reset_vector_database(persist_directory, embeddings):
    if os.path.exists(persist_directory):
        try:
            existing_vectordb = Chroma(persist_directory=persist_directory, embedding=embeddings)
            if len(existing_vectordb) > 0:
                print(f"Existing Chroma DB contains {len(existing_vectordb)} entries.")
                user_input = input("Reset DB? (y/n): ")
                if user_input.lower() != "y":
                    print("Skipping reset.")
                    return
        except Exception:
            pass
        print(f"Cleaning up old DB at {persist_directory}...")
        shutil.rmtree(persist_directory)
        print("Old DB cleaned up.")

# Setup Paths
persist_directory = "./chroma_db2"
data_dir = "./data/pdf"
pdf_pattern = os.path.join(data_dir, "*.pdf")
docx_pattern = os.path.join(data_dir, "*.docx")

# Initialize Embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-mpnet-base-v2")

# Reset Vector DB
reset_vector_database(persist_directory, embeddings)

# Load PDF & DOCX Files
pdf_files = glob.glob(pdf_pattern)
docx_files = glob.glob(docx_pattern)

if not pdf_files and not docx_files:
    raise ValueError(f"No PDF or DOCX files found in {data_dir}")

print(f"📂 Found {len(pdf_files)} PDFs and {len(docx_files)} Word documents in {data_dir}.")

# Preprocess and Vectorize
documents = preprocess_documents(pdf_files, docx_files)
vectordb = Chroma.from_documents(documents=documents, embedding=embeddings, persist_directory=persist_directory)

print(f"✅ Stored {len(documents)} documents in ChromaDB at '{persist_directory}'.")
